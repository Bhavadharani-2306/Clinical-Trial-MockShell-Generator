import tempfile
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import re
import logging

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


def _set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'bottom', 'left', 'right', 'insideH', 'insideV'):
        tag = OxmlElement(f'w:{edge}')
        attrs = kwargs.get(edge, {"val": "nil"})
        for k, v in attrs.items():
            tag.set(qn(f'w:{k}'), str(v))
        tcBorders.append(tag)
    tcPr.append(tcBorders)


def _set_table_no_borders(table):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for edge in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        tag = OxmlElement(f'w:{edge}')
        tag.set(qn('w:val'), 'nil')
        tblBorders.append(tag)
    tblPr.append(tblBorders)


def _add_bottom_border_to_cell(cell, size=6, color="000000"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ['top', 'left', 'right', 'insideH', 'insideV']:
        tag = OxmlElement(f'w:{edge}')
        tag.set(qn('w:val'), 'nil')
        tcBorders.append(tag)
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(size))
    bottom.set(qn('w:color'), color)
    tcBorders.append(bottom)
    tcPr.append(tcBorders)


def _add_solid_horizontal_rule(doc, size='6', color='000000'):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(size))
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


class Exporter:
    def __init__(self, output_dir=None):
        self.output_dir = output_dir or tempfile.gettempdir()
        if not os.path.exists(self.output_dir):
            os.makedirs(self.output_dir)

    def _normalize_and_clean_templates(self, raw_templates):
        return raw_templates if raw_templates else []

    def _sort_templates(self, templates):
        return self._normalize_and_clean_templates(templates)

    def export_to_word(self, templates, filename="MockShells.docx"):
        doc = Document()
        sorted_shells = self._sort_templates(templates)
        
        section = doc.sections[0]
        section.orientation = 1
        section.page_width, section.page_height = Inches(11), Inches(8.5)
        section.left_margin = section.right_margin = Inches(0.75)

        style = doc.styles['Normal']
        style.font.name = 'Courier New'
        style.font.size = Pt(9)

        # Title Cover
        for _ in range(3): doc.add_paragraph()
        p_title = doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_title.add_run("CLINICAL TRIAL STATISTICAL MOCK SHELLS\n").font.size = Pt(18)
        doc.add_page_break()

        for shell in sorted_shells:
            p = doc.add_paragraph()
            p.add_run(f"{shell.get('type')} {shell.get('number')} : {shell.get('title')}").font.bold = True
            _add_solid_horizontal_rule(doc, size=8)
            
            headers = shell.get("headers", [])
            rows = shell.get("rows", [])
            if headers:
                table = doc.add_table(rows=1, cols=len(headers))
                _set_table_no_borders(table)
                for i, h_text in enumerate(headers):
                    cell = table.rows[0].cells[i]
                    cell.paragraphs[0].add_run(str(h_text)).font.bold = True
                    _add_bottom_border_to_cell(cell, size=6)
                for row_data in rows:
                    row_cells = table.add_row().cells
                    for i, val in enumerate(row_data):
                        if i < len(row_cells):
                            row_cells[i].paragraphs[0].add_run(str(val))
                            _set_cell_border(row_cells[i])
            _add_solid_horizontal_rule(doc, size=6)
            doc.add_paragraph(shell.get('notes_section', ''))
            doc.add_paragraph(shell.get('prog_notes', ''))
            doc.add_page_break()

        path = os.path.join(self.output_dir, filename)
        doc.save(path)
        return path

    def sanitize_pdf_text(self, text):
        if not text: return ""
        mapping = {'\u2022': '-', '\u2013': '-', '\u2014': '-', '\u201c': '"', '\u201d': '"', '\u2018': "'", '\u2019': "'"}
        for uc, rep in mapping.items():
            text = str(text).replace(uc, rep)
        return text.encode('latin-1', 'replace').decode('latin-1')

    def export_to_pdf(self, templates, filename="MockShells.pdf"):
        """Creates pharma-standard PDF mock shell document with Title, Reviewers and TOC."""
        try:
            from fpdf import FPDF
            sorted_shells = self._sort_templates(templates)

            if not sorted_shells:
                sorted_shells = [{
                    "type": "Table", "number": "X.X", "title": "No Layout Artifacts Compiled", 
                    "company": "Global Pharma Inc.", "protocol": "PROTOCOL-XYZ-123", "headers": [], "rows": []
                }]

            class MockPDF(FPDF):
                def header(self): pass
                def footer(self):
                    self.set_y(-12)
                    self.set_font('Courier', 'I', 7)
                    self.cell(0, 5, f'Page {self.page_no()}', 0, 0, 'R')

            pdf = MockPDF(orientation='L', unit='mm', format='Letter')
            pdf.set_auto_page_break(auto=True, margin=18)
            pdf.set_margins(15, 15, 15)

            sample = sorted_shells[0] if sorted_shells else {}
            company  = self.sanitize_pdf_text(sample.get('company', 'Global Pharma Inc.'))
            protocol = self.sanitize_pdf_text(sample.get('protocol', 'PROTOCOL-XYZ-123'))

            # ── 1. TITLE PAGE ─────────────────────────────────────────────────
            pdf.add_page()
            pdf.set_font('Courier', 'B', 18)
            pdf.ln(40)
            pdf.cell(0, 10, "CLINICAL TRIAL STATISTICAL MOCK SHELLS", 0, 1, 'C')
            pdf.set_font('Courier', 'I', 12)
            pdf.cell(0, 8, f"Protocol Specification Package: {protocol}", 0, 1, 'C')
            pdf.ln(30)
            pdf.set_font('Courier', '', 10)
            pdf.multi_cell(0, 6, f"Sponsor: {company}\nDate of Generation: 2026\nClassification: Regulatory Confidential", 0, 'C')

            # ── 2. REVIEWER PANEL SIGN-OFF MATRIX ─────────────────────────────────
            pdf.add_page()
            pdf.set_font('Courier', 'B', 12)
            pdf.cell(0, 8, "REGULATORY REVIEW & APPROVAL SIGN-OFF PANEL", 0, 1, 'L')
            pdf.set_font('Courier', '', 9)
            pdf.cell(0, 6, "The technical layout specifications detailed within this manual have been reviewed and approved:", 0, 1, 'L')
            pdf.ln(4)
            
            # Print reviewer table grid stubs
            pdf.set_font('Courier', 'B', 8)
            pdf.cell(50, 6, "Functional Role", 1, 0, 'L')
            pdf.cell(70, 6, "Reviewer Name / Title", 1, 0, 'L')
            pdf.cell(80, 6, "Signature Stub", 1, 0, 'L')
            pdf.cell(40, 6, "Date", 1, 1, 'L')
            pdf.set_font('Courier', '', 8)
            for r in ["Lead Biostatistician", "Statistical Programmer", "Clinical Data Manager"]:
                pdf.cell(50, 8, r, 1, 0, 'L')
                pdf.cell(70, 8, "", 1, 0, 'L')
                pdf.cell(80, 8, "", 1, 0, 'L')
                pdf.cell(40, 8, "", 1, 1, 'L')

            # ── 3. TABLE OF CONTENTS ──────────────────────────────────────────────
            pdf.add_page()
            pdf.set_font('Courier', 'B', 12)
            pdf.cell(0, 8, "TABLE OF CONTENTS (AUTOMATED MANIFEST)", 0, 1, 'L')
            pdf.set_font('Courier', '', 9)
            pdf.cell(0, 6, "Planned Output Artifacts Summary List:", 0, 1, 'L')
            pdf.ln(2)
            
            pdf.set_font('Courier', '', 8)
            for shell in sorted_shells:
                t_type = shell.get('type', 'Table')
                t_num = shell.get('number', 'X.X')
                t_title = self.sanitize_pdf_text(shell.get('title', 'Untitled'))
                pdf.cell(0, 5, f"- {t_type} {t_num} : {t_title}", 0, 1, 'L')

            # ── 4. SHELLS PARSING MATRIX ──────────────────────────────────────
            for shell in sorted_shells:
                pdf.add_page()
                tlf_type = self.sanitize_pdf_text(shell.get('type', 'Table'))
                tlf_num  = self.sanitize_pdf_text(shell.get('number', 'X.X.X'))
                title    = self.sanitize_pdf_text(shell.get('title', 'Untitled'))
                pop      = self.sanitize_pdf_text(shell.get('population', 'Intent-to-Treat Population'))

                # Page Header
                pdf.set_font('Courier', 'B', 8)
                pdf.cell(140, 5, f"{company}    Protocol: {protocol}", 0, 0, 'L')
                pdf.cell(0, 5, 'DRAFT / CONFIDENTIAL', 0, 1, 'R')
                pdf.ln(2)

                # Unified Title Identification Block
                pdf.set_font('Courier', 'B', 9)
                pdf.multi_cell(0, 5, f"{tlf_type} {tlf_num} : {title}", 0, 'L')
                
                # Dark Solid Underline Rule
                pdf.set_draw_color(0, 0, 0)
                pdf.set_line_width(0.4)
                pdf.line(pdf.l_margin, pdf.get_y() + 1, pdf.l_margin + 247, pdf.get_y() + 1)
                pdf.ln(3)

                # Population Field
                pdf.set_font('Courier', 'I', 8)
                pdf.cell(0, 5, f"Analysis Population: {pop}", 0, 1, 'L')
                pdf.ln(2)

                # Structured Table Rendering Grid
                headers = shell.get("headers", [])
                rows    = shell.get("rows", [])

                if headers:
                    n_cols = len(headers)
                    page_w = 247
                    label_w = int(page_w * 0.30)
                    stat_w  = int(page_w * 0.10)
                    rem     = page_w - label_w - stat_w
                    data_w  = rem // max(n_cols - 2, 1) if n_cols > 2 else rem

                    col_widths = [label_w, stat_w] + [data_w] * max(n_cols - 2, 0)

                    pdf.set_font('Courier', 'B', 7)
                    for i, h in enumerate(headers):
                        san_h = self.sanitize_pdf_text(str(h)).replace('\n', ' ')
                        w = col_widths[i] if i < len(col_widths) else data_w
                        align = 'L' if i == 0 else 'C'
                        pdf.cell(w, 6, san_h, 0, 0, align)
                    pdf.ln()
                    
                    pdf.line(pdf.l_margin, pdf.get_y(), pdf.l_margin + page_w, pdf.get_y())
                    pdf.ln(1)

                    pdf.set_font('Courier', '', 7)
                    for row_data in rows:
                        for i, val in enumerate(row_data):
                            if i >= n_cols: break
                            w = col_widths[i] if i < len(col_widths) else data_w
                            san_val = self.sanitize_pdf_text(str(val))[:35]
                            align = 'L' if i == 0 else 'C'
                            pdf.cell(w, 5, san_val, 0, 0, align)
                        pdf.ln()

                # Footer Unified Separation Line
                pdf.ln(2)
                pdf.set_draw_color(0, 0, 0)
                pdf.set_line_width(0.4)
                pdf.line(pdf.l_margin, pdf.get_y(), pdf.l_margin + 247, pdf.get_y())
                pdf.ln(2)

                # Notes and Metadata block sections
                pdf.set_font('Courier', '', 7)
                notes = self.sanitize_pdf_text(shell.get('notes_section', ''))
                if notes:
                    pdf.multi_cell(0, 4, notes, 0, 'L')
                    pdf.ln(1)

                footnotes = self.sanitize_pdf_text(shell.get('footnotes', ''))
                if footnotes and footnotes != "Details to be finalized":
                    pdf.multi_cell(0, 4, f"[1] {footnotes}", 0, 'L')
                    pdf.ln(1)

                prog = self.sanitize_pdf_text(shell.get('prog_notes', ''))
                if prog:
                    pdf.set_font('Courier', 'I', 7)
                    pdf.multi_cell(0, 4, prog, 0, 'L')

            path = os.path.join(self.output_dir, filename)
            pdf.output(path)
            return path
        except Exception as e:
            logger.error(f"PDF Export Failed: {str(e)}", exc_info=True)
            return None

    def export_to_excel(self, templates, filename="MockShells.xlsx"):
        """Generates dynamic multi-tab clinical layouts in Excel using openpyxl."""
        try:
            import openpyxl
            from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
            from openpyxl.utils import get_column_letter

            wb = openpyxl.Workbook()
            # Remove default worksheet
            default_sheet = wb.active
            wb.remove(default_sheet)

            sorted_shells = self._sort_templates(templates)

            for shell in sorted_shells:
                t_type = shell.get('type', 'Table')
                t_num = shell.get('number', 'X.X')
                t_title = shell.get('title', 'Untitled')
                
                # Excel worksheet names are limited to 31 characters
                sheet_name = f"{t_type[:5]} {t_num}"[:31]
                for char in r"\/?:*[]":
                    sheet_name = sheet_name.replace(char, "_")
                
                ws = wb.create_sheet(title=sheet_name)
                ws.views.sheetView[0].showGridLines = True

                # Formatted Professional Typography
                title_font = Font(name="Calibri", size=14, bold=True, color="0f766e")
                meta_font = Font(name="Calibri", size=10, italic=True)
                header_font = Font(name="Calibri", size=11, bold=True, color="ffffff")
                header_fill = PatternFill(start_color="0f766e", end_color="0f766e", fill_type="solid")
                data_font = Font(name="Calibri", size=11)
                footnote_font = Font(name="Calibri", size=9, italic=True)

                thin_border = Border(
                    left=Side(style='thin', color='cbd5e1'),
                    right=Side(style='thin', color='cbd5e1'),
                    top=Side(style='thin', color='cbd5e1'),
                    bottom=Side(style='thin', color='cbd5e1')
                )

                # Top Block Metrics
                current_row = 1
                ws.cell(row=current_row, column=1, value=f"{shell.get('company', 'Global Pharma Inc.')} | Protocol: {shell.get('protocol', 'PROTOCOL-XYZ-123')}").font = meta_font
                current_row += 1

                ws.cell(row=current_row, column=1, value=f"{t_type} {t_num} : {t_title}").font = title_font
                current_row += 1

                pop_text = f"Analysis Population: {shell.get('population', 'Safety Population')}"
                ws.cell(row=current_row, column=1, value=pop_text).font = meta_font
                current_row += 2 # Clean visual buffer

                # Draw Table Header Layout Grid
                headers = shell.get("headers", [])
                rows = shell.get("rows", [])

                if headers:
                    for col_idx, h in enumerate(headers, 1):
                        cell = ws.cell(row=current_row, column=col_idx, value=str(h))
                        cell.font = header_font
                        cell.fill = header_fill
                        cell.alignment = Alignment(horizontal="center" if col_idx > 1 else "left", vertical="center", wrap_text=True)
                        cell.border = thin_border
                    ws.row_dimensions[current_row].height = 28
                    current_row += 1

                    # Write Raw Placeholder Row Data
                    for row_data in rows:
                        for col_idx, val in enumerate(row_data, 1):
                            cell = ws.cell(row=current_row, column=col_idx, value=str(val))
                            cell.font = data_font
                            cell.alignment = Alignment(horizontal="center" if col_idx > 1 else "left", vertical="center")
                            cell.border = thin_border
                        ws.row_dimensions[current_row].height = 20
                        current_row += 1

                current_row += 1 # Grid offset spacing

                # Append Clinical Footnotes and Source Paths
                notes = shell.get('notes_section', '')
                if notes:
                    ws.cell(row=current_row, column=1, value=notes).font = footnote_font
                    current_row += 1

                footnotes = shell.get('footnotes', '')
                if footnotes:
                    for line in footnotes.split('\n'):
                        if line.strip():
                            ws.cell(row=current_row, column=1, value=line.strip()).font = footnote_font
                            current_row += 1

                prog_path = shell.get('validated_source_code_path', '')
                if prog_path:
                    ws.cell(row=current_row, column=1, value=f"Source SAS Path: {prog_path}").font = footnote_font
                    current_row += 1

                # Auto-Adjust dynamic column widths
                for col in ws.columns:
                    max_len = 0
                    col_letter = get_column_letter(col[0].column)
                    for cell in col:
                        if cell.value:
                            # Only parse row ranges within the table data grid to prevent long titles from stretching col A
                            if 5 <= cell.row < (5 + len(rows)):
                                max_len = max(max_len, len(str(cell.value)))
                    ws.column_dimensions[col_letter].width = max(max_len + 5, 15)

            path = os.path.join(self.output_dir, filename)
            wb.save(path)
            return path
        except Exception as e:
            logger.error(f"Excel Export Failed: {str(e)}", exc_info=True)
            return None

    def export_to_txt(self, templates, filename="MockShells.txt"):
        sorted_shells = self._sort_templates(templates)
        path = os.path.join(self.output_dir, filename)
        with open(path, "w", encoding="utf-8") as f:
            for shell in sorted_shells:
                f.write(f"{shell.get('type')} {shell.get('number')} : {shell.get('title')}\n")
                f.write("=" * 80 + "\n")
                hdrs = shell.get('headers', [])
                if hdrs:
                    f.write(" | ".join(str(h).replace('\n', ' ') for h in hdrs) + "\n")
                    f.write("-" * 80 + "\n")
                    for r in shell.get('rows', []):
                        f.write(" | ".join(str(v).replace('\n', ' ') for v in r) + "\n")
                f.write("\n" + shell.get('notes_section', '') + "\n" + shell.get('prog_notes', '') + "\n\n")
        return path

    def export_to_md(self, templates, filename="MockShells.md"):
        sorted_shells = self._sort_templates(templates)
        path = os.path.join(self.output_dir, filename)
        with open(path, "w", encoding="utf-8") as f:
            for shell in sorted_shells:
                f.write(f"## {shell.get('type')} {shell.get('number')}: {shell.get('title')}\n\n")
                hdrs = shell.get('headers', [])
                if hdrs:
                    f.write("| " + " | ".join(str(h).replace('\n', '<br>') for h in hdrs) + " |\n")
                    f.write("| " + " | ".join("---" for _ in hdrs) + " |\n")
                    for r in shell.get('rows', []):
                        f.write("| " + " | ".join(str(v).replace('\n', '<br>') for v in r) + " |\n")
                f.write(f"\n```text\n{shell.get('notes_section')}\n{shell.get('prog_notes')}\n```\n\n---\n\n")
        return path